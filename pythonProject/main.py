from docx import Document
from docx.shared import Inches
import xlsxwriter

authors = ['Иванов Иван Иванович', 'Петров Петр Петрович']

for x in authors:
    document = Document()
    document.add_heading('Визитка', 0)
    p = document.add_paragraph(x + '\n менеджер')
    p.add_run('\n консультация в области информационных технологий')
    p.add_run('\n Тел.: 8 777 777 77 77' + '\n e-mail: python@mail.ru')
    document.add_heading('Первая консультация бесплатная', level=1)
    document.add_paragraph('')
    document.add_picture('1.png', width=Inches(2))
    document.add_page_break()
    document.save(x + '.docx')

try:
    my_file = 'Визитка.xlsx'
    book = xlsxwriter.Workbook(my_file)
    sheet = book.add_worksheet()
    sheet.set_column('A:A', 25)
    sheet.set_column('B:B', 2.2)

    bold = book.add_format({'bold': True})
    sheet.write('C1', 'Иванов Иван Иванович', bold)
    sheet.write('C2', 'менеджер')
    sheet.write('C3', 'консультация в области информационных технологий')
    sheet.write(3, 2, 'Тел.: 8 777 777 77 77')
    sheet.write(4, 2, 'e-mail: python@mail.ru')

    sheet.insert_image('A1', '1.png')
    book.close()
except Exception as a:
    print("Error!")
    print(a)

if __name__ == '_main_':
    main()

